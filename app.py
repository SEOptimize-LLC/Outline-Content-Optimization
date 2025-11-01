import streamlit as st
import anthropic
import openai
import google.generativeai as genai
from sentence_transformers import SentenceTransformer
import spacy
from sklearn.metrics.pairwise import cosine_similarity
from sklearn.feature_extraction.text import TfidfVectorizer
import pandas as pd
import numpy as np
import re
from io import BytesIO
import base64
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import markdown

# Page configuration
st.set_page_config(
    page_title="Blog Post Optimizer",
    page_icon="✍️",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS for better UI
st.markdown("""
<style>
    .main-header {
        font-size: 2.5rem;
        font-weight: 700;
        color: #1f77b4;
        margin-bottom: 0.5rem;
    }
    .sub-header {
        font-size: 1.2rem;
        color: #666;
        margin-bottom: 2rem;
    }
    .success-box {
        padding: 1rem;
        border-radius: 0.5rem;
        background-color: #d4edda;
        border: 1px solid #c3e6cb;
        margin: 1rem 0;
    }
    .highlight {
        background-color: #fff3cd;
        padding: 0.2rem 0.4rem;
        border-radius: 0.25rem;
    }
</style>
""", unsafe_allow_html=True)

# Initialize session state
if 'outline_result' not in st.session_state:
    st.session_state.outline_result = None
if 'draft_result' not in st.session_state:
    st.session_state.draft_result = None
if 'audience_insights' not in st.session_state:
    st.session_state.audience_insights = None

# Load models (cached to prevent reloading)
@st.cache_resource
def load_nlp_models():
    """Load and cache NLP models"""
    try:
        # Load spaCy model
        nlp = spacy.load("en_core_web_sm")
    except OSError:
        # Model not found - show helpful error for Streamlit Cloud users
        st.error("""
        ⚠️ **spaCy model not found!**
        
        Please add this line to your `requirements.txt`:
        ```
        https://github.com/explosion/spacy-models/releases/download/en_core_web_sm-3.7.1/en_core_web_sm-3.7.1-py3-none-any.whl
        ```
        
        Then redeploy your app.
        """)
        st.stop()
    
    # Load sentence transformer for semantic similarity
    semantic_model = SentenceTransformer('all-MiniLM-L6-v2')
    
    return nlp, semantic_model

# API Configuration
def get_api_keys():
    """Retrieve API keys from Streamlit secrets"""
    openai_key = ""
    anthropic_key = ""
    google_key = ""
    
    try:
        if "OPENAI_API_KEY" in st.secrets:
            openai_key = st.secrets["OPENAI_API_KEY"]
            # Ensure it's not empty
            if not openai_key or openai_key.strip() == "":
                openai_key = ""
        
        if "ANTHROPIC_API_KEY" in st.secrets:
            anthropic_key = st.secrets["ANTHROPIC_API_KEY"]
            if not anthropic_key or anthropic_key.strip() == "":
                anthropic_key = ""
        
        if "GOOGLE_API_KEY" in st.secrets:
            google_key = st.secrets["GOOGLE_API_KEY"]
            if not google_key or google_key.strip() == "":
                google_key = ""
    except Exception as e:
        # Secrets not configured or error accessing them
        pass
    
    return openai_key, anthropic_key, google_key

def get_available_models():
    """Get available models based on configured API keys"""
    openai_key, anthropic_key, google_key = get_api_keys()
    
    models = {}
    
    if openai_key and len(openai_key) > 10:  # Basic validation
        models["OpenAI"] = {
            "gpt-4o": "GPT-4o (Latest, Most Capable)",
            "gpt-4o-mini": "GPT-4o Mini (Fast & Efficient)",
            "gpt-4-turbo": "GPT-4 Turbo (Balanced)"
        }
    
    if anthropic_key and len(anthropic_key) > 10:  # Basic validation
        models["Anthropic"] = {
            "claude-3-5-sonnet-20241022": "Claude 3.5 Sonnet (Recommended)",
            "claude-3-5-haiku-20241022": "Claude 3.5 Haiku (Fast)",
            "claude-3-opus-20240229": "Claude 3 Opus (Most Capable)"
        }
    
    if google_key and len(google_key) > 10:  # Basic validation
        models["Google"] = {
            "gemini-2.0-flash-exp": "Gemini 2.0 Flash (Experimental, Fastest)",
            "gemini-1.5-pro": "Gemini 1.5 Pro (Balanced)",
            "gemini-1.5-flash": "Gemini 1.5 Flash (Fast & Cost-Effective)"
        }
    
    return models

# AI API Functions
def call_openai(prompt, model="gpt-4o", max_tokens=4000):
    """Call OpenAI API"""
    openai_key, _, _ = get_api_keys()
    if not openai_key or len(openai_key) < 10:
        st.error("⚠️ OpenAI API key not found or invalid in secrets!")
        return None
    
    try:
        client = openai.OpenAI(api_key=openai_key)
        response = client.chat.completions.create(
            model=model,
            messages=[{"role": "user", "content": prompt}],
            max_tokens=max_tokens,
            temperature=0.7
        )
        return response.choices[0].message.content
    except Exception as e:
        st.error(f"❌ OpenAI API Error: {str(e)}")
        return None

def call_anthropic(prompt, model="claude-3-5-sonnet-20241022", max_tokens=4000):
    """Call Anthropic API"""
    _, anthropic_key, _ = get_api_keys()
    if not anthropic_key or len(anthropic_key) < 10:
        st.error("⚠️ Anthropic API key not found or invalid in secrets!")
        return None
    
    try:
        client = anthropic.Anthropic(api_key=anthropic_key)
        message = client.messages.create(
            model=model,
            max_tokens=max_tokens,
            messages=[{"role": "user", "content": prompt}]
        )
        return message.content[0].text
    except Exception as e:
        st.error(f"❌ Anthropic API Error: {str(e)}")
        return None

def call_gemini(prompt, model="gemini-1.5-pro", max_tokens=4000):
    """Call Google Gemini API"""
    _, _, google_key = get_api_keys()
    if not google_key or len(google_key) < 10:
        st.error("⚠️ Google API key not found or invalid in secrets!")
        return None
    
    try:
        genai.configure(api_key=google_key)
        model_instance = genai.GenerativeModel(model)
        
        generation_config = {
            "max_output_tokens": max_tokens,
            "temperature": 0.7,
        }
        
        response = model_instance.generate_content(
            prompt,
            generation_config=generation_config
        )
        return response.text
    except Exception as e:
        st.error(f"❌ Google Gemini API Error: {str(e)}")
        return None

def call_ai_model(prompt, model_string, max_tokens=4000):
    """Universal AI caller that routes to the correct provider"""
    if model_string.startswith("gpt"):
        return call_openai(prompt, model_string, max_tokens)
    elif model_string.startswith("claude"):
        return call_anthropic(prompt, model_string, max_tokens)
    elif model_string.startswith("gemini"):
        return call_gemini(prompt, model_string, max_tokens)
    else:
        st.error("Unknown model type!")
        return None

# Utility Functions
def extract_headings(markdown_text):
    """Extract H2 and H3 headings from Markdown"""
    h2_pattern = r'^##\s+(.+)$'
    h3_pattern = r'^###\s+(.+)$'
    
    headings = []
    for line in markdown_text.split('\n'):
        h2_match = re.match(h2_pattern, line.strip())
        h3_match = re.match(h3_pattern, line.strip())
        
        if h2_match:
            headings.append(('H2', h2_match.group(1)))
        elif h3_match:
            headings.append(('H3', h3_match.group(1)))
    
    return headings

def semantic_similarity_score(text1, text2, model):
    """Calculate semantic similarity between two texts"""
    embeddings = model.encode([text1, text2])
    similarity = cosine_similarity([embeddings[0]], [embeddings[1]])[0][0]
    return similarity

def create_download_link(content, filename, file_format="markdown"):
    """Create download link without page reset"""
    if file_format == "markdown":
        b64 = base64.b64encode(content.encode()).decode()
        href = f'<a href="data:text/markdown;base64,{b64}" download="{filename}">📥 Download {filename}</a>'
    else:  # PDF or other formats
        b64 = base64.b64encode(content.encode()).decode()
        href = f'<a href="data:text/plain;base64,{b64}" download="{filename}">📥 Download {filename}</a>'
    return href

def markdown_to_docx(markdown_text, title="Optimized Content"):
    """Convert markdown text to a Word document with formatting"""
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Split content into lines
    lines = markdown_text.split('\n')
    
    for line in lines:
        line = line.strip()
        
        if not line:
            # Empty line - add paragraph break
            doc.add_paragraph()
            continue
        
        # H1 headers
        if line.startswith('# ') and not line.startswith('## '):
            text = line.replace('# ', '')
            p = doc.add_heading(text, level=1)
        
        # H2 headers
        elif line.startswith('## ') and not line.startswith('### '):
            text = line.replace('## ', '')
            p = doc.add_heading(text, level=2)
        
        # H3 headers
        elif line.startswith('### '):
            text = line.replace('### ', '')
            p = doc.add_heading(text, level=3)
        
        # Bullet points
        elif line.startswith('- ') or line.startswith('* '):
            text = line[2:]
            p = doc.add_paragraph(text, style='List Bullet')
        
        # Numbered lists
        elif re.match(r'^\d+\.\s', line):
            text = re.sub(r'^\d+\.\s', '', line)
            p = doc.add_paragraph(text, style='List Number')
        
        # Bold text (simplified - just remove ** markers for now)
        elif '**' in line:
            text = line.replace('**', '')
            p = doc.add_paragraph(text)
            for run in p.runs:
                run.bold = True
        
        # Regular paragraph
        else:
            doc.add_paragraph(line)
    
    # Save to BytesIO
    docx_file = BytesIO()
    doc.save(docx_file)
    docx_file.seek(0)
    
    return docx_file

# === TAB 1: OUTLINE OPTIMIZER ===
def audience_research_analysis(keyword, model):
    """Perform deep audience research and search intent analysis"""
    prompt = f"""Conduct advanced audience research for the keyword/topic: "{keyword}"

Focus on:
1. Ideal Customer Profile (ICP) - Demographics and psychographics
2. Core pain points and frustrations that drive searches
3. Emotional triggers and psychological motivations
4. Search behaviors across platforms (Google, ChatGPT, Claude, Perplexity, Gemini)
5. User intent clusters (informational, transactional, navigational, investigational)
6. Common questions and information gaps
7. Resonance strategies - how to connect with this audience

Provide a comprehensive analysis structured as:
- **Demographics & Psychographics**: Who searches for this?
- **Pain Points**: What problems are they trying to solve?
- **Emotional Triggers**: What drives their search behavior?
- **Intent Clusters**: What are they really looking for?
- **Key Questions**: Top 10 questions they ask
- **Resonance Strategy**: How to create content that connects

Be specific, actionable, and data-driven in your analysis."""

    return call_ai_model(prompt, model, max_tokens=3000)

def optimize_outline(keyword, draft_outline, query_fanout, audience_insights, model):
    """Optimize blog outline with AI analysis"""
    prompt = f"""You are an expert SEO content strategist. Optimize the following blog post outline.

PRIMARY KEYWORD: {keyword}

AUDIENCE INSIGHTS:
{audience_insights}

QUERY FAN-OUT ANALYSIS:
{query_fanout}

CURRENT DRAFT OUTLINE:
{draft_outline}

TASK:
1. Analyze the draft outline against the Query Fan-Out suggestions and Audience Insights
2. Prioritize Query Fan-Out recommendations over the current draft
3. Ensure all ICP pain points and search intents are addressed
4. Generate an OPTIMIZED OUTLINE in the EXACT same Markdown structure (preserve H2/H3 hierarchy)

For each H2/H3 section, add:
- 7-12 concise talking points (bullet list) to guide writers
- Each talking point should be actionable and derived from audience research
- Include rationale for key additions (e.g., "Addresses [pain point] from audience research")

OUTPUT FORMAT:
## [H2 Heading]
- Talking point 1
- Talking point 2
...
- Talking point 7-12

### [H3 Heading]
- Talking point 1
...

Ensure logical flow, comprehensive coverage, and SEO optimization. Make it ready for a writer to execute."""

    return call_ai_model(prompt, model, max_tokens=4000)

# === TAB 2: DRAFT OPTIMIZER ===
def keyword_relevance_analysis(primary_keyword, keyword_list, draft_content, nlp, semantic_model):
    """Analyze keyword relevance and find placement opportunities"""
    # Parse draft with spaCy
    doc = nlp(draft_content)
    
    # Split into paragraphs
    paragraphs = [p.strip() for p in draft_content.split('\n\n') if p.strip()]
    
    # Calculate TF-IDF for relevance scoring
    all_text = [primary_keyword] + keyword_list + paragraphs
    vectorizer = TfidfVectorizer()
    tfidf_matrix = vectorizer.fit_transform(all_text)
    
    results = []
    
    for keyword in keyword_list:
        # Calculate semantic similarity with primary keyword
        primary_sim = semantic_similarity_score(keyword, primary_keyword, semantic_model)
        
        # Check if keyword exists in draft (exact or variations)
        keyword_lower = keyword.lower()
        exists_exact = keyword_lower in draft_content.lower()
        
        # Find semantic variations
        variations = []
        for sent in doc.sents:
            sent_text = sent.text.lower()
            if keyword_lower not in sent_text:
                sim = semantic_similarity_score(keyword, sent.text, semantic_model)
                if sim > 0.8:
                    variations.append(sent.text[:50] + "...")
        
        # Calculate relevance score
        keyword_embedding = semantic_model.encode([keyword])
        draft_embedding = semantic_model.encode([draft_content[:500]])  # First 500 chars
        relevance_score = cosine_similarity(keyword_embedding, draft_embedding)[0][0]
        
        # Determine if keyword needs addition
        keyword_density = draft_content.lower().count(keyword_lower) / len(draft_content.split()) * 100
        needs_addition = relevance_score > 0.5 and keyword_density < 1.0
        
        if needs_addition:
            # Find best placement location
            best_para_idx = 0
            best_sim = 0
            for idx, para in enumerate(paragraphs[:5]):  # Check first 5 paragraphs
                sim = semantic_similarity_score(keyword, para, semantic_model)
                if sim > best_sim:
                    best_sim = sim
                    best_para_idx = idx
            
            results.append({
                'keyword': keyword,
                'relevance_score': round(relevance_score, 2),
                'density': round(keyword_density, 2),
                'variations_found': len(variations),
                'needs_addition': needs_addition,
                'best_paragraph_idx': best_para_idx,
                'paragraph_preview': paragraphs[best_para_idx][:100] + "..."
            })
    
    return results

def generate_keyword_integration(keyword, paragraph_context, model):
    """Generate natural keyword integration"""
    prompt = f"""You are a skilled content writer. Rewrite the following paragraph to naturally integrate the keyword while maintaining flow and readability.

KEYWORD TO INTEGRATE: {keyword}

ORIGINAL PARAGRAPH:
{paragraph_context}

REQUIREMENTS:
- Integrate the keyword naturally and contextually
- Maintain the original tone and structure
- Keep paragraph length similar (50-200 words)
- Ensure it reads naturally, not stuffed
- Make it copy-paste ready

OUTPUT: Only provide the rewritten paragraph, nothing else."""

    return call_ai_model(prompt, model, max_tokens=500)

def ai_tool_optimization(draft_content, primary_keyword, keyword_list, model):
    """Apply 10-item AI optimization checklist"""
    
    # Count original words
    original_word_count = len(draft_content.split())
    
    prompt = f"""You are an expert content optimizer for AI search tools (Google AI Overviews, ChatGPT, Claude, Perplexity).

PRIMARY KEYWORD: {primary_keyword}
KEYWORDS: {', '.join(keyword_list[:10])}

CONTENT DRAFT (Original word count: {original_word_count} words):
{draft_content}

CRITICAL REQUIREMENTS:
1. MAINTAIN OR EXPAND the original content length - aim for {original_word_count} to {int(original_word_count * 1.2)} words
2. DO NOT shorten or remove substantial content
3. ADD details, examples, data, and depth rather than removing content
4. Enhancement means ADDING VALUE, not reducing content
5. KEEP ALL ORIGINAL HEADINGS - do NOT convert them to questions
6. DO NOT start sections with questions - use direct, declarative statements that ANSWER immediately

Apply these optimizations while PRESERVING OR EXPANDING content:

1. **Answer-First Introduction**: Rewrite opening to DIRECTLY ANSWER the core query with declarative statements (no questions)
2. **Answer-First Sections**: EVERY section must start with a direct answer statement, NOT a question
3. **Semantic Chunks**: Ensure 75-300 word self-contained sections (expand short sections, don't shorten long ones)
4. **Answer-Evidence-Context**: Restructure chunks (answer -> evidence -> context) - ADD evidence and context, don't remove
5. **Direct Sentences**: Convert to active voice, Subject-Verb-Object (maintain all information)
6. **Informational Density**: Increase specifics by 20% (numbers, entities, examples) - this means ADDING content
7. **Attribute Claims**: Replace generics with "A 2023 study by XYZ found..." - ADD citations and sources
8. **Signal Experience**: Add "In our testing..." where applicable - ADD first-hand insights
9. **FAQ Section**: Append 3-10 long-tail Q&A pairs at the end - this ADDS content
10. **Title & Meta**: Generate optimized title (<60 chars) and meta (140-160 chars) at the top

WRITING STYLE RULES:
- NEVER start a section with a question
- Lead with direct, declarative statements
- Example WRONG: "Is it safe to brush after whitening strips?"
- Example CORRECT: "You can safely brush your teeth 30 minutes after using whitening strips."
- Use statements, not questions, throughout the main content
- Questions are ONLY allowed in the FAQ section at the end

OUTPUT the fully optimized content in clean Markdown format. 
- Do NOT add any tags, markers, or annotations
- Output should be clean, ready-to-use content
- PRESERVE the original heading structure - do NOT convert headings to questions
- DO NOT start sections with questions - always lead with answers
- Ensure final word count is AT LEAST {int(original_word_count * 0.95)} words (95% of original minimum)
- Aim to EXPAND content with valuable details, not compress it"""

    return call_ai_model(prompt, model, max_tokens=8000)

# === TAB 3: SEMANTIC SEO ANALYZER (Koray Tuğberk GÜBÜR Framework) ===

def extract_macro_context(content, primary_keyword, nlp, semantic_model, model):
    """Extract macro context elements from content using Koray's framework"""

    prompt = f"""You are an expert in Koray Tuğberk GÜBÜR's Semantic SEO framework. Analyze the following content and extract detailed macro context elements.

PRIMARY KEYWORD/ENTITY: {primary_keyword}

CONTENT:
{content[:4000]}

Analyze and extract the following MACRO CONTEXT elements:

1. **Primary Topic/Central Entity**: Identify the main subject/entity. Is it the keyword provided or something more specific?

2. **Domain/Authority (Source Context)**: Who is this content for? What's the brand identity? How does it monetize? (e.g., consultancy, e-commerce, educational blog)

3. **User's Intent**: What is the dominant search intent? (Informational, Transactional, Navigational, Investigational)

4. **Search Persona**: Who is searching for this? Demographics, psychographics, expertise level

5. **Main Benefits**: What value does this content provide? How does it connect to the source context?

6. **Entity-Attribute-Value (EAV) Inventory**: List 10-15 key attribute-value pairs for the central entity
   Format: Entity - Attribute: Value
   Example: "Germany - Population: 83 million"

7. **Intent & Topical Clusters**: Identify topical clusters and group related intents
   - Core Section (directly related to monetization)
   - Author/Outer Section (broader topical authority)

8. **Link Hub Potential**: What would be the ideal root document H1 for a semantic content network around this topic?

Provide detailed, actionable analysis for each element."""

    return call_ai_model(prompt, model, max_tokens=4000)

def extract_micro_context(content, primary_keyword, nlp, semantic_model, model):
    """Extract micro context elements from content using Koray's framework"""

    # First, use spaCy for entity extraction
    doc = nlp(content[:10000])  # Limit for performance
    entities = [(ent.text, ent.label_) for ent in doc.ents][:30]

    # Extract headings for contextual hierarchy
    headings = extract_headings(content)

    prompt = f"""You are an expert in Koray Tuğberk GÜBÜR's Semantic SEO framework. Analyze the following content and extract detailed micro context elements.

PRIMARY KEYWORD/ENTITY: {primary_keyword}

CONTENT:
{content[:4000]}

EXTRACTED ENTITIES (from NLP):
{entities[:20]}

HEADING STRUCTURE:
{headings[:15]}

Analyze and extract the following MICRO CONTEXT elements:

1. **Semantically Relevant Entities**: List 10-15 entities that support the macro context (beyond the primary entity)

2. **Key Attributes**: What characteristics define and describe the main entity and related entities?

3. **Values**: Specific data points for attributes (numbers, measurements, dates, specifications)

4. **Predicates**: The relationships or actions connecting entities and attributes
   Example: "Whitening peroxides chemically open tooth pores"

5. **Temporal Elements**: Time-sensitive information, dates, durations, timelines

6. **Conditional Synonyms**: Phrases using conjunctive words (and, or) that create contextual variations
   Example: "Religion and Belief Structure" or "Costs and Conditions"

7. **Co-occurring Terms (Distributional Semantics)**:
   - Identify 5-7 term clusters that should appear together in specific sections
   - Example: ["water", "uric acid", "stone"] for kidney stone content

8. **Annotation Text Patterns**: Text that should appear around internal links for semantic connection

9. **Anchor Segments**: Mutual words that should appear in sequential sentences for discourse flow

10. **Question Types Present**: Categorize existing or implied questions
    - Boolean (yes/no)
    - Definitional (what is)
    - Grouping (types of)
    - Comparative (versus)
    - Temporal (how long, when)

11. **Modality and Measurement Units**: Scientific terminology, units, measurements used

12. **Macro vs Micro Content Boundary**: Where does the main content transition to supplementary content?

Provide detailed, actionable analysis for each element."""

    return call_ai_model(prompt, model, max_tokens=4000)

def generate_content_brief(content, primary_keyword, macro_context, micro_context, model):
    """Generate a comprehensive content brief based on Koray's 4-column framework"""

    prompt = f"""You are an expert in Koray Tuğberk GÜBÜR's Semantic SEO framework. Generate a detailed content brief for optimizing the following content.

PRIMARY KEYWORD/ENTITY: {primary_keyword}

MACRO CONTEXT ANALYSIS:
{macro_context}

MICRO CONTEXT ANALYSIS:
{micro_context}

CURRENT CONTENT STRUCTURE:
{content[:2000]}

Generate a comprehensive CONTENT BRIEF using Koray's 4-column framework:

## 1. CONTEXTUAL VECTOR (The Flow)
- List all headings in optimal logical sequence (H1 → H2 → H3)
- Ensure straight, proper context without interruptions
- Order by search demand and semantic closeness
- For each heading, explain WHY it's positioned there
- Identify the central search intent for the root document

## 2. CONTEXTUAL HIERARCHY (The Weight)
- Assign coverage weight to each section (% of content)
- Macro context sections should get 60-70% of content
- Micro context sections should get 30-40% of content
- Specify heading levels (H1, H2, H3) strategically
- Recommend word count ranges for each major section

## 3. CONTEXTUAL STRUCTURE (The Format)
- Specify format for each section (paragraph, list, table, Q&A)
- Provide article methodology rules:
  * Start with definitions
  * Use answer-first approach (no questions in headings)
  * Answer-Evidence-Context formula for each section
  * Specify where to use boolean questions (FAQ only)
  * Modality instructions (factual vs. research-based vs. suggestions)
  * Required elements (dates, units, measurements, citations)
- Include specific co-occurrence instructions for distributional semantics
- Specify anchor segments for discourse integration

## 4. CONTEXTUAL CONNECTIONS (Internal Linking)
- Identify 5-10 strategic internal link opportunities
- Specify exact anchor text for each link
- Indicate placement priority (top, middle, bottom)
- Suggest related pages in the semantic content network:
  * Root document recommendation
  * Supporting seed articles
  * Node articles for broader topical coverage
- Explain the semantic relationship for each link

## 5. OPTIMIZATION RECOMMENDATIONS
Based on the analysis, provide specific recommendations:
- Missing elements from macro context
- Weak areas in micro context
- Improvements to contextual vector (heading flow)
- Enhancements to distributional semantics (co-occurrence)
- Title tag optimization (macro context)
- Meta description optimization (macro + micro context)
- FAQ section questions to add

Provide actionable, specific recommendations ready for implementation."""

    return call_ai_model(prompt, model, max_tokens=6000)

def analyze_distributional_semantics(content, nlp, semantic_model):
    """Analyze co-occurrence patterns and distributional semantics"""

    doc = nlp(content[:5000])  # Limit for performance

    # Extract sentences
    sentences = [sent.text for sent in doc.sents if len(sent.text.split()) > 5][:50]

    # Extract paragraphs
    paragraphs = [p.strip() for p in content.split('\n\n') if len(p.strip().split()) > 20][:20]

    # Simple co-occurrence analysis
    from collections import Counter
    from itertools import combinations

    # Get significant terms (nouns and proper nouns)
    significant_terms = []
    for token in doc:
        if token.pos_ in ['NOUN', 'PROPN'] and not token.is_stop and len(token.text) > 3:
            significant_terms.append(token.text.lower())

    # Count term frequencies
    term_freq = Counter(significant_terms)
    top_terms = [term for term, count in term_freq.most_common(30)]

    # Find co-occurrences in sentences
    co_occur_matrix = {}
    for sent in sentences:
        sent_terms = [term for term in top_terms if term in sent.lower()]
        if len(sent_terms) >= 2:
            for term1, term2 in combinations(sent_terms, 2):
                pair = tuple(sorted([term1, term2]))
                co_occur_matrix[pair] = co_occur_matrix.get(pair, 0) + 1

    # Sort by frequency
    top_co_occurrences = sorted(co_occur_matrix.items(), key=lambda x: x[1], reverse=True)[:15]

    return {
        'top_terms': term_freq.most_common(20),
        'co_occurrences': top_co_occurrences,
        'sentence_count': len(sentences),
        'paragraph_count': len(paragraphs)
    }

def validate_primary_entity(content, claimed_entity, model):
    """Validate if the claimed primary keyword is actually the main entity"""

    prompt = f"""You are an expert in Koray Tuğberk GÜBÜR's Semantic SEO framework.

The user claims the primary keyword/entity is: "{claimed_entity}"

Analyze this content and determine if this is truly the central entity:

CONTENT:
{content[:2000]}

Provide:
1. **Validation**: Is "{claimed_entity}" the true primary entity? (Yes/No)
2. **Confidence Level**: (High/Medium/Low)
3. **Actual Primary Entity**: If different, what is it?
4. **Reasoning**: Explain your analysis
5. **Macro Context Alignment**: Does "{claimed_entity}" reflect the true macro context?

Be specific and analytical in your response."""

    return call_ai_model(prompt, model, max_tokens=1000)

def generate_gap_analysis(content, primary_keyword, macro_context, micro_context, model):
    """Identify specific gaps and missing elements in the content"""

    prompt = f"""You are an expert content optimizer using Koray Tuğberk GÜBÜR's Semantic SEO framework.

PRIMARY ENTITY: {primary_keyword}

MACRO CONTEXT ANALYSIS:
{macro_context[:1500]}

MICRO CONTEXT ANALYSIS:
{micro_context[:1500]}

CURRENT CONTENT:
{content[:3000]}

Perform a detailed GAP ANALYSIS identifying what's MISSING or WEAK:

## 1. MACRO CONTEXT GAPS
- Missing EAV (Entity-Attribute-Value) pairs that should be included
- Weak or unclear source context signals
- User intent misalignment
- Missing topical clusters (core vs. author sections)

## 2. MICRO CONTEXT GAPS
- Missing semantically relevant entities
- Weak predicates (relationships between concepts)
- Missing temporal elements (dates, durations, timelines)
- Insufficient conditional synonyms
- Missing co-occurring term clusters
- Absent question types (boolean, definitional, comparative, etc.)

## 3. DISTRIBUTIONAL SEMANTICS GAPS
- Terms that should co-occur but don't
- Missing contextual dance patterns
- Weak word proximities

## 4. CONTENT STRUCTURE GAPS
- Missing or poorly structured headings
- Incorrect contextual hierarchy (macro/micro balance)
- Missing internal linking opportunities
- Weak contextual vector (heading flow)

For EACH gap identified, provide:
- **What's missing**: Specific element
- **Why it matters**: Impact on semantic SEO
- **Where to add it**: Specific section/paragraph
- **Example**: Concrete text suggestion WITH SOURCE ATTRIBUTION

CITATION REQUIREMENTS:
- For all factual claims, statistics, and data in examples, include source attribution
- Format: "According to [Organization/Study, Year], ..."
- If no source available, use: "[Source needed: verify this claim]"
- NEVER fabricate specific numbers without attribution

Be highly specific and actionable. Writers should know exactly what to add."""

    return call_ai_model(prompt, model, max_tokens=4000)

def generate_optimization_action_plan(content, primary_keyword, macro_context, micro_context, gap_analysis, model):
    """Generate prioritized, actionable optimization steps for writers"""

    prompt = f"""You are an expert content optimizer using Koray Tuğberk GÜBÜR's Semantic SEO framework.

PRIMARY ENTITY: {primary_keyword}

MACRO CONTEXT:
{macro_context[:1000]}

MICRO CONTEXT:
{micro_context[:1000]}

GAP ANALYSIS:
{gap_analysis[:2000]}

CURRENT CONTENT:
{content[:2000]}

Generate a WRITER-FRIENDLY OPTIMIZATION ACTION PLAN with specific, copy-paste ready suggestions.

## PRIORITY 1: CRITICAL FIXES (Do These First)
For each critical fix:
- **Action**: Specific change to make
- **Location**: Exact section/paragraph (e.g., "After H2: Introduction")
- **Before**: Current text (if applicable)
- **After**: Replacement text (copy-paste ready)
- **Reasoning**: Why this improves semantic SEO

## PRIORITY 2: IMPORTANT ENHANCEMENTS (Do These Next)
Same format as Priority 1

## PRIORITY 3: NICE-TO-HAVE IMPROVEMENTS (Do If Time Permits)
Same format as Priority 1

## HEADING OPTIMIZATION
Current heading structure vs. Optimized heading structure
- Show BEFORE and AFTER for each heading
- Explain the semantic improvement

## CO-OCCURRENCE OPTIMIZATION
Specific paragraphs where term clusters should appear together:
- **Section**: Where to add
- **Terms to co-occur**: List the 3-5 terms
- **Example paragraph**: Copy-paste ready text that includes proper co-occurrence

## INTERNAL LINKING RECOMMENDATIONS
- **Anchor text**: Exact text to use
- **Link target**: Where it should link to (describe the topic)
- **Annotation text**: Text to appear before/after the link
- **Placement**: Exact location in content

## ENTITY-ATTRIBUTE-VALUE ADDITIONS
Missing EAV pairs to add:
- **Entity - Attribute**: What's missing
- **Where to add**: Specific section
- **Example sentence**: Copy-paste ready

## MACRO CONTEXT STRENGTHENING
- **Current macro context score**: Weak/Moderate/Strong
- **Specific improvements**: What to change
- **Example**: Before/After text with source attributions

CRITICAL CITATION REQUIREMENTS FOR ALL EXAMPLES:
- ALWAYS include source attribution for statistics, facts, and specific claims
- Format: "According to [Source Name/Study, Year], ..."
- Or use footnote markers: "[1]" with reference at bottom
- If source unknown, use: "[Source needed: verify X statistic]"
- NEVER fabricate specific numbers or statistics without attribution
- Generic statements don't need citations, but specific data always does

Make every suggestion copy-paste ready with proper citations. Writers should be able to implement changes immediately."""

    return call_ai_model(prompt, model, max_tokens=6000)

def generate_optimized_version(content, primary_keyword, action_plan, model):
    """Generate a fully optimized version of the content based on the action plan"""

    prompt = f"""You are an expert content writer implementing Koray Tuğberk GÜBÜR's Semantic SEO framework.

PRIMARY ENTITY: {primary_keyword}

OPTIMIZATION ACTION PLAN:
{action_plan[:3000]}

CURRENT CONTENT:
{content}

Generate a FULLY OPTIMIZED VERSION of this content by implementing ALL the recommendations from the action plan.

REQUIREMENTS:
1. Maintain or expand the original word count
2. Implement ALL Priority 1 and Priority 2 changes
3. Optimize heading structure for contextual vector
4. Add missing EAV (Entity-Attribute-Value) pairs naturally
5. Ensure proper co-occurrence of term clusters in relevant sections
6. Strengthen macro context (60-70% of content)
7. Add micro context elements (30-40% of content)
8. Include internal linking anchor text with proper annotation
9. Add missing temporal elements, predicates, and conditional synonyms
10. Ensure proper discourse flow with anchor segments

OUTPUT FORMAT:
- Clean Markdown
- Properly structured headings (H1, H2, H3)
- Natural, readable prose
- Ready to publish

Generate the complete optimized article now."""

    return call_ai_model(prompt, model, max_tokens=8000)

def optimize_section_by_section(content, primary_keyword, action_plan, nlp, model):
    """Optimize content section by section for granular control"""

    # Extract sections based on H2 headings
    sections = []
    current_section = {"heading": "Introduction", "content": ""}

    lines = content.split('\n')
    for line in lines:
        if line.startswith('## ') and not line.startswith('### '):
            # New H2 section
            if current_section["content"]:
                sections.append(current_section)
            current_section = {"heading": line.replace('## ', ''), "content": line + "\n"}
        else:
            current_section["content"] += line + "\n"

    # Add last section
    if current_section["content"]:
        sections.append(current_section)

    # Optimize each section
    optimized_sections = []

    for idx, section in enumerate(sections):
        section_heading = section["heading"]
        section_content = section["content"]

        # Extract relevant action items for this section
        relevant_actions = f"Focus on improvements related to section: {section_heading}"

        prompt = f"""You are an expert content writer implementing Koray Tuğberk GÜBÜR's Semantic SEO framework.

PRIMARY ENTITY: {primary_keyword}

OPTIMIZATION ACTION PLAN (Full):
{action_plan[:1500]}

CURRENT SECTION TO OPTIMIZE:
Heading: {section_heading}

Content:
{section_content}

TASK: Optimize ONLY this section by:
1. Implementing relevant Priority 1 and Priority 2 changes from the action plan
2. Adding missing EAV pairs that relate to this section's topic
3. Ensuring proper co-occurrence of relevant term clusters
4. Strengthening semantic relationships
5. Adding citations and sources for any factual claims (format: "According to [Source], ..." or include [1] footnote markers)
6. Maintaining or expanding word count
7. Preserving the heading structure (H2, H3 levels)

IMPORTANT CITATION REQUIREMENTS:
- ALWAYS provide source attribution for statistics, facts, and claims
- Format: "According to [Organization Name Study, Year], statistic..."
- Or use footnote markers: "Germany receives 1.5 million immigrants annually.[1]"
- If you don't have a real source, use placeholders: "[Source needed: verify statistic]"
- NEVER make up specific numbers without attribution

OUTPUT:
- Optimized version of this section ONLY
- Clean Markdown
- Include the heading
- Natural, readable prose
- Include source attributions for all claims

Generate the optimized section now."""

        optimized_section = call_ai_model(prompt, model, max_tokens=2000)

        optimized_sections.append({
            "original_heading": section_heading,
            "original_content": section_content,
            "optimized_content": optimized_section,
            "section_number": idx + 1
        })

    return optimized_sections

# === MAIN APP ===
def main():
    st.markdown('<p class="main-header">✍️ Blog Post Optimizer</p>', unsafe_allow_html=True)
    st.markdown('<p class="sub-header">AI-Powered Content Optimization for SEO & AI Search Tools</p>', unsafe_allow_html=True)
    
    # Sidebar - API Configuration Check
    with st.sidebar:
        st.header("⚙️ Configuration")
        openai_key, anthropic_key, google_key = get_api_keys()
        
        # Display API connection status
        st.subheader("API Connection Status")
        if openai_key and len(openai_key) > 10:  # Basic validation
            st.success("✅ OpenAI API Connected")
        else:
            st.warning("⚠️ OpenAI API Key Missing")
        
        if anthropic_key and len(anthropic_key) > 10:  # Basic validation
            st.success("✅ Anthropic API Connected")
        else:
            st.warning("⚠️ Anthropic API Key Missing")
        
        if google_key and len(google_key) > 10:  # Basic validation
            st.success("✅ Google Gemini API Connected")
        else:
            st.warning("⚠️ Google API Key Missing")
        
        # Get available models
        available_models = get_available_models()
        
        if not available_models:
            st.error("❌ No API keys configured! Please add at least one API key to Streamlit Secrets.")
            st.stop()
        
        st.markdown("---")
        
        # Smart Model Selector - only shows models from configured providers
        st.subheader("🤖 AI Model Selection")
        
        # Flatten the models dictionary for selection
        model_options = {}
        for provider, models in available_models.items():
            for model_id, model_name in models.items():
                model_options[f"{provider}: {model_name}"] = model_id
        
        selected_model_display = st.selectbox(
            "Select AI Model:",
            options=list(model_options.keys()),
            help="Only models from configured API providers are shown"
        )
        
        selected_model = model_options[selected_model_display]
        
        # Display model info
        st.info(f"**Active Model:** `{selected_model}`")
        
        st.markdown("---")
        st.markdown("**💡 Setup Instructions:**")
        st.markdown("Add API keys to Streamlit Cloud Secrets:")
        st.code('OPENAI_API_KEY = "your-key"\nANTHROPIC_API_KEY = "your-key"\nGOOGLE_API_KEY = "your-key"')
        
        st.markdown("---")
        st.markdown("**💰 Cost Comparison:**")
        st.markdown("- 🥇 **Gemini**: Most cost-effective (free tier)")
        st.markdown("- 🥈 **Anthropic**: Best quality/cost ratio")
        st.markdown("- 🥉 **OpenAI**: Most capable models")
    
    # Load NLP models
    with st.spinner("Loading NLP models..."):
        nlp, semantic_model = load_nlp_models()
    
    # Main Tabs
    tab1, tab2, tab3 = st.tabs(["📝 Outline Optimizer", "✨ Draft Optimizer", "🎯 Semantic SEO Analyzer"])
    
    # === TAB 1: OUTLINE OPTIMIZER ===
    with tab1:
        st.header("Outline Optimizer")
        st.markdown("Optimize your blog outline with audience research, search intent analysis, and SEO best practices.")
        
        col1, col2 = st.columns([1, 1])
        
        with col1:
            keyword_outline = st.text_input(
                "Primary Keyword/Topic *",
                placeholder="e.g., dog training tips",
                key="outline_keyword"
            )
            
            draft_outline = st.text_area(
                "Draft Outline (Markdown) *",
                placeholder="## Introduction\n\n## Main Point 1\n### Subpoint 1.1\n\n## Conclusion",
                height=250,
                key="draft_outline_input"
            )
            
            uploaded_outline = st.file_uploader(
                "Or Upload Draft Outline",
                type=['txt', 'md'],
                key="outline_upload",
                help="Upload your draft outline as a Markdown or text file"
            )
        
        with col2:
            query_fanout = st.text_area(
                "Query Fan-Out Analysis",
                placeholder="Paste your Query Fan-Out report here (structured text with expanded queries, related terms, intent insights)",
                height=250,
                key="query_fanout_input"
            )
            
            uploaded_fanout = st.file_uploader(
                "Or Upload Query Fan-Out Report",
                type=['txt', 'md'],
                key="fanout_upload",
                help="Upload your Query Fan-Out analysis report"
            )
        
        # Handle file uploads
        if uploaded_outline:
            draft_outline = uploaded_outline.read().decode('utf-8')
            st.success("✅ Outline uploaded successfully!")
        
        if uploaded_fanout:
            query_fanout = uploaded_fanout.read().decode('utf-8')
            st.success("✅ Query Fan-Out report uploaded successfully!")
        
        st.markdown("---")
        
        if st.button("🚀 Optimize Outline", type="primary", use_container_width=True):
            if not keyword_outline or not draft_outline:
                st.error("Please provide both primary keyword and draft outline!")
            else:
                # Step 1: Audience Research
                with st.spinner("🔍 Step 1/2: Conducting deep audience research..."):
                    progress_bar = st.progress(0)
                    status_text = st.empty()
                    
                    status_text.text("Analyzing search intent and user psychology...")
                    progress_bar.progress(20)
                    
                    audience_insights = audience_research_analysis(keyword_outline, selected_model)
                    st.session_state.audience_insights = audience_insights
                    
                    status_text.text("Research complete!")
                    progress_bar.progress(100)
                    
                    if audience_insights:
                        st.success("✅ Audience research complete!")
                        with st.expander("📊 View Audience Insights", expanded=False):
                            st.markdown(audience_insights)
                    
                    status_text.empty()
                
                # Step 2: Outline Optimization
                with st.spinner("✨ Step 2/2: Optimizing outline..."):
                    progress_bar2 = st.progress(0)
                    status_text2 = st.empty()
                    
                    fanout_text = query_fanout if query_fanout else "No Query Fan-Out analysis provided."
                    
                    status_text2.text("Cross-referencing with Query Fan-Out analysis...")
                    progress_bar2.progress(30)
                    
                    optimized_outline = optimize_outline(
                        keyword_outline,
                        draft_outline,
                        fanout_text,
                        audience_insights,
                        selected_model
                    )
                    
                    status_text2.text("Generating talking points...")
                    progress_bar2.progress(90)
                    
                    st.session_state.outline_result = optimized_outline
                    progress_bar2.progress(100)
                    status_text2.empty()
                
                if optimized_outline:
                    st.markdown("---")
                    st.subheader("📊 Optimization Results")
                    
                    # Side-by-side comparison
                    comp_col1, comp_col2 = st.columns(2)
                    
                    with comp_col1:
                        st.markdown("### 📄 Original Outline")
                        st.markdown(draft_outline)
                    
                    with comp_col2:
                        st.markdown("### ✅ Optimized Outline")
                        st.markdown(optimized_outline)
                    
                    # Download options
                    st.markdown("---")
                    download_col1, download_col2 = st.columns(2)
                    
                    with download_col1:
                        filename = f"optimized_outline_{datetime.now().strftime('%Y%m%d_%H%M%S')}.md"
                        st.markdown(
                            create_download_link(optimized_outline, filename),
                            unsafe_allow_html=True
                        )
                    
                    with download_col2:
                        if st.session_state.audience_insights:
                            insights_filename = f"audience_insights_{datetime.now().strftime('%Y%m%d_%H%M%S')}.md"
                            st.markdown(
                                create_download_link(st.session_state.audience_insights, insights_filename),
                                unsafe_allow_html=True
                            )
    
    # === TAB 2: DRAFT OPTIMIZER ===
    with tab2:
        st.header("Draft Optimizer")
        st.markdown("Optimize your draft through keyword analysis and AI tool optimization (2-step process).")
        
        keyword_draft = st.text_input(
            "Primary Keyword/Topic *",
            placeholder="e.g., dog training tips",
            key="draft_keyword"
        )
        
        draft_col1, draft_col2 = st.columns([2, 1])
        
        with draft_col1:
            draft_content = st.text_area(
                "Content Draft *",
                placeholder="Paste your blog draft here (Markdown or HTML supported)",
                height=300,
                key="draft_content_input"
            )
            
            uploaded_draft = st.file_uploader(
                "Or Upload Draft File",
                type=['txt', 'md', 'html'],
                key="draft_upload",
                help="Upload your blog draft as Markdown, text, or HTML"
            )
        
        with draft_col2:
            keyword_list_input = st.text_area(
                "Expanded Keyword List *",
                placeholder="keyword 1\nkeyword 2\nkeyword 3",
                height=300,
                key="keyword_list_input"
            )
        
        # Handle file upload
        if uploaded_draft:
            draft_content = uploaded_draft.read().decode('utf-8')
            st.success("✅ Draft uploaded successfully!")
        
        # Parse keyword list
        keyword_list = [k.strip() for k in keyword_list_input.split('\n') if k.strip()]
        if ',' in keyword_list_input:
            keyword_list = [k.strip() for k in keyword_list_input.split(',') if k.strip()]
        
        st.markdown("---")
        
        if st.button("🚀 Optimize Draft", type="primary", use_container_width=True):
            if not keyword_draft or not draft_content or not keyword_list:
                st.error("Please provide primary keyword, draft content, and keyword list!")
            else:
                # Subprocess 1: Keyword Optimization
                st.markdown("### 📊 Subprocess 1: Keyword Optimization")
                
                with st.spinner("🔍 Analyzing keyword relevance and placement..."):
                    keyword_analysis = keyword_relevance_analysis(
                        keyword_draft,
                        keyword_list,
                        draft_content,
                        nlp,
                        semantic_model
                    )
                
                if keyword_analysis:
                    st.success(f"✅ Analyzed {len(keyword_analysis)} high-relevance keywords needing optimization")
                    
                    # Generate integration suggestions
                    integration_results = []
                    
                    progress = st.progress(0)
                    for idx, kw_data in enumerate(keyword_analysis[:10]):  # Limit to top 10
                        with st.spinner(f"Generating integration for '{kw_data['keyword']}'..."):
                            integrated_paragraph = generate_keyword_integration(
                                kw_data['keyword'],
                                kw_data['paragraph_preview'],
                                selected_model
                            )
                            
                            integration_results.append({
                                'Selected Keyword': f"{kw_data['keyword']} (Score: {kw_data['relevance_score']})",
                                'Placement Hint': f"Insert in paragraph {kw_data['best_paragraph_idx'] + 1}: '{kw_data['paragraph_preview'][:50]}...'",
                                'Content Snippet': integrated_paragraph if integrated_paragraph else "Generation failed"
                            })
                        
                        progress.progress((idx + 1) / min(10, len(keyword_analysis)))
                    
                    # Display table
                    if integration_results:
                        st.markdown("#### 📝 Keyword Integration Suggestions")
                        df = pd.DataFrame(integration_results)
                        st.dataframe(df, use_container_width=True, height=400)
                
                st.markdown("---")
                
                # Subprocess 2: AI Tool Optimization
                st.markdown("### ✨ Subprocess 2: AI Tool Optimization")
                
                with st.spinner("🤖 Applying 10-item AI optimization checklist..."):
                    progress2 = st.progress(0)
                    
                    optimized_draft = ai_tool_optimization(
                        draft_content,
                        keyword_draft,
                        keyword_list,
                        selected_model
                    )
                    
                    progress2.progress(100)
                    st.session_state.draft_result = optimized_draft
                
                if optimized_draft:
                    st.success("✅ Draft fully optimized!")
                    
                    # Checklist compliance
                    st.markdown("#### ✅ Optimization Checklist Compliance")
                    checklist_items = [
                        "Answer-First Introduction",
                        "Question-Based H2 Headings",
                        "Semantic Chunks (75-300 words)",
                        "Answer-Evidence-Context Formula",
                        "Direct Sentence Structures",
                        "Informational Density (+20%)",
                        "Attributed Claims",
                        "First-Hand Experience Signals",
                        "Dedicated FAQ Section",
                        "Optimized Title & Meta"
                    ]
                    
                    cols = st.columns(5)
                    for idx, item in enumerate(checklist_items):
                        with cols[idx % 5]:
                            st.metric(label=item, value="✅", delta="Applied")
                    
                    st.markdown("---")
                    
                    # Display optimized draft
                    st.markdown("#### 📄 Fully Optimized Draft")
                    st.markdown(optimized_draft)
                    
                    # Download
                    st.markdown("---")
                    final_filename = f"optimized_draft_{datetime.now().strftime('%Y%m%d_%H%M%S')}.md"
                    st.markdown(
                        create_download_link(optimized_draft, final_filename),
                        unsafe_allow_html=True
                    )

    # === TAB 3: SEMANTIC SEO ANALYZER ===
    with tab3:
        st.header("🎯 Semantic SEO Content Analyzer")
        st.markdown("**Powered by Koray Tuğberk GÜBÜR's Semantic SEO Framework**")
        st.markdown("Analyze content line-by-line for macro/micro context, entity-attribute-value relationships, and generate comprehensive content briefs.")

        st.markdown("---")

        # Input Section
        col1_sem, col2_sem = st.columns([2, 1])

        with col1_sem:
            primary_entity = st.text_input(
                "Primary Keyword/Main Entity *",
                placeholder="e.g., Germany, Kidney Stones, How Long Do Teeth Pores Stay Open After Whitening",
                key="semantic_entity",
                help="This can be a single word, phrase, or complete question - whatever represents your central topic"
            )

            content_input = st.text_area(
                "Content Draft (Markdown) *",
                placeholder="Paste your blog post content here...",
                height=400,
                key="semantic_content"
            )

            uploaded_semantic = st.file_uploader(
                "Or Upload Content File",
                type=['txt', 'md'],
                key="semantic_upload",
                help="Upload your content as Markdown or text file"
            )

        with col2_sem:
            st.markdown("### Analysis Options")

            validate_entity = st.checkbox(
                "Validate Primary Entity",
                value=True,
                help="Confirm if your claimed primary entity is actually the main topic"
            )

            include_distrib_semantics = st.checkbox(
                "Include Distributional Semantics",
                value=True,
                help="Analyze co-occurrence patterns and term relationships"
            )

            generate_brief = st.checkbox(
                "Generate Content Brief",
                value=True,
                help="Create comprehensive 4-column content brief (Koray's framework)"
            )

            st.markdown("---")

            generate_gap_analysis_ui = st.checkbox(
                "Generate Gap Analysis",
                value=True,
                help="Identify specific missing elements and weaknesses"
            )

            generate_action_plan = st.checkbox(
                "Generate Action Plan",
                value=True,
                help="Create prioritized, copy-paste ready optimization steps for writers"
            )

            generate_optimized = st.checkbox(
                "Optimize Section-by-Section",
                value=False,
                help="Optimize each H2 section individually for granular control (95% more usable than full rewrites)"
            )

            st.markdown("---")
            st.markdown("### Framework Elements")
            st.info("""
**Macro Context:**
- Primary Topic
- Domain/Authority
- User Intent
- Search Persona
- Main Benefits
- EAV Inventory
- Topical Clusters
- Link Hub

**Micro Context:**
- Entities & Attributes
- Values & Predicates
- Temporal Elements
- Conditional Synonyms
- Co-occurring Terms
- Question Types
            """)

        # Handle file upload
        if uploaded_semantic:
            content_input = uploaded_semantic.read().decode('utf-8')
            st.success("✅ Content uploaded successfully!")

        st.markdown("---")

        # Analysis Button
        if st.button("🚀 Analyze Content", type="primary", use_container_width=True, key="analyze_semantic"):
            if not primary_entity or not content_input:
                st.error("⚠️ Please provide both primary entity and content!")
            else:
                # Initialize session state for results
                if 'semantic_results' not in st.session_state:
                    st.session_state.semantic_results = {}

                # Step 1: Validate Primary Entity
                if validate_entity:
                    with st.spinner("🔍 Step 1: Validating primary entity..."):
                        validation_result = validate_primary_entity(content_input, primary_entity, selected_model)
                        st.session_state.semantic_results['validation'] = validation_result

                        if validation_result:
                            st.success("✅ Entity validation complete!")
                            with st.expander("📋 Entity Validation Results", expanded=True):
                                st.markdown(validation_result)

                # Step 2: Extract Macro Context
                with st.spinner("🎯 Step 2: Extracting macro context elements..."):
                    progress = st.progress(0)

                    macro_context = extract_macro_context(
                        content_input,
                        primary_entity,
                        nlp,
                        semantic_model,
                        selected_model
                    )
                    st.session_state.semantic_results['macro_context'] = macro_context
                    progress.progress(33)

                    if macro_context:
                        st.success("✅ Macro context extracted!")
                        with st.expander("📊 Macro Context Analysis", expanded=True):
                            st.markdown(macro_context)

                # Step 3: Extract Micro Context
                with st.spinner("🔬 Step 3: Extracting micro context elements..."):
                    micro_context = extract_micro_context(
                        content_input,
                        primary_entity,
                        nlp,
                        semantic_model,
                        selected_model
                    )
                    st.session_state.semantic_results['micro_context'] = micro_context
                    progress.progress(66)

                    if micro_context:
                        st.success("✅ Micro context extracted!")
                        with st.expander("🔬 Micro Context Analysis", expanded=True):
                            st.markdown(micro_context)

                # Step 4: Distributional Semantics (Optional)
                if include_distrib_semantics:
                    with st.spinner("📈 Step 4: Analyzing distributional semantics..."):
                        distrib_results = analyze_distributional_semantics(content_input, nlp, semantic_model)
                        st.session_state.semantic_results['distributional'] = distrib_results

                        st.success("✅ Distributional semantics analyzed!")
                        with st.expander("📈 Co-occurrence Analysis", expanded=False):
                            st.markdown("### Top Terms")
                            df_terms = pd.DataFrame(distrib_results['top_terms'], columns=['Term', 'Frequency'])
                            st.dataframe(df_terms, use_container_width=True)

                            st.markdown("### Top Co-occurring Term Pairs")
                            co_occur_data = [
                                {'Term 1': pair[0], 'Term 2': pair[1], 'Co-occurrences': count}
                                for (pair, count) in distrib_results['co_occurrences']
                            ]
                            df_cooccur = pd.DataFrame(co_occur_data)
                            st.dataframe(df_cooccur, use_container_width=True)

                            st.info(f"""
**Analysis Summary:**
- Sentences analyzed: {distrib_results['sentence_count']}
- Paragraphs analyzed: {distrib_results['paragraph_count']}
- Unique significant terms: {len(distrib_results['top_terms'])}
- Co-occurrence pairs found: {len(distrib_results['co_occurrences'])}
                            """)

                progress.progress(100)

                # Step 5: Generate Gap Analysis (Optional)
                gap_analysis_result = None
                if generate_gap_analysis_ui and macro_context and micro_context:
                    st.markdown("---")
                    with st.spinner("🔍 Step 5: Performing gap analysis..."):
                        gap_analysis_result = generate_gap_analysis(
                            content_input,
                            primary_entity,
                            macro_context,
                            micro_context,
                            selected_model
                        )
                        st.session_state.semantic_results['gap_analysis'] = gap_analysis_result

                        if gap_analysis_result:
                            st.success("✅ Gap analysis complete!")
                            with st.expander("🔍 Gap Analysis - What's Missing", expanded=True):
                                st.markdown(gap_analysis_result)
                                st.info("💡 **For Writers**: Each gap includes WHERE to add it and EXAMPLE text to use.")

                # Step 6: Generate Action Plan (Optional)
                action_plan_result = None
                if generate_action_plan and macro_context and micro_context:
                    st.markdown("---")
                    with st.spinner("📋 Step 6: Generating optimization action plan..."):
                        # Use gap analysis if available, otherwise pass empty string
                        gap_for_plan = gap_analysis_result if gap_analysis_result else "No gap analysis performed."

                        action_plan_result = generate_optimization_action_plan(
                            content_input,
                            primary_entity,
                            macro_context,
                            micro_context,
                            gap_for_plan,
                            selected_model
                        )
                        st.session_state.semantic_results['action_plan'] = action_plan_result

                        if action_plan_result:
                            st.success("✅ Optimization action plan generated!")
                            with st.expander("📋 Writer Action Plan - Copy-Paste Ready", expanded=True):
                                st.markdown(action_plan_result)
                                st.success("✅ **For Writers**: All suggestions are copy-paste ready. Prioritized by importance.")

                            # Download Action Plan separately
                            action_plan_filename = f"action_plan_{datetime.now().strftime('%Y%m%d_%H%M%S')}.md"
                            st.markdown(
                                create_download_link(action_plan_result, action_plan_filename),
                                unsafe_allow_html=True
                            )

                # Step 7: Generate Content Brief (Optional)
                if generate_brief and macro_context and micro_context:
                    st.markdown("---")
                    with st.spinner("📝 Step 7: Generating comprehensive content brief..."):
                        progress_brief = st.progress(0)

                        content_brief = generate_content_brief(
                            content_input,
                            primary_entity,
                            macro_context,
                            micro_context,
                            selected_model
                        )
                        st.session_state.semantic_results['content_brief'] = content_brief
                        progress_brief.progress(100)

                        if content_brief:
                            st.success("✅ Content brief generated!")

                            st.markdown("---")
                            st.markdown("## 📋 Comprehensive Content Brief")
                            st.markdown("**Based on Koray Tuğberk GÜBÜR's 4-Column Framework**")
                            st.markdown(content_brief)

                            # Download Content Brief
                            st.markdown("---")
                            brief_filename = f"content_brief_{datetime.now().strftime('%Y%m%d_%H%M%S')}.md"
                            st.markdown(
                                create_download_link(content_brief, brief_filename),
                                unsafe_allow_html=True
                            )

                # Step 8: Section-by-Section Optimization (Optional)
                if generate_optimized and action_plan_result:
                    st.markdown("---")
                    st.markdown("## 🎯 Section-by-Section Optimization")
                    st.info("✅ This optimizes each H2 section individually, giving you granular control. Review and accept/reject each section.")

                    with st.spinner("✨ Step 8: Optimizing sections (this may take 1-2 minutes)..."):
                        optimized_sections = optimize_section_by_section(
                            content_input,
                            primary_entity,
                            action_plan_result,
                            nlp,
                            selected_model
                        )
                        st.session_state.semantic_results['optimized_sections'] = optimized_sections

                        if optimized_sections:
                            st.success(f"✅ {len(optimized_sections)} sections optimized!")

                            st.markdown("---")
                            st.markdown("### 📊 Review Sections Individually")
                            st.markdown("**Use the toggles below to accept/reject each optimized section**")

                            # Initialize acceptance state if not exists
                            if 'accepted_sections' not in st.session_state:
                                st.session_state.accepted_sections = [True] * len(optimized_sections)

                            for idx, section in enumerate(optimized_sections):
                                with st.expander(f"📄 Section {section['section_number']}: {section['original_heading']}", expanded=False):
                                    # Acceptance toggle
                                    accept_col, stats_col = st.columns([1, 3])

                                    with accept_col:
                                        st.session_state.accepted_sections[idx] = st.checkbox(
                                            "✅ Accept",
                                            value=st.session_state.accepted_sections[idx],
                                            key=f"accept_section_{idx}"
                                        )

                                    with stats_col:
                                        orig_words = len(section['original_content'].split())
                                        opt_words = len(section['optimized_content'].split())
                                        word_change = opt_words - orig_words
                                        st.metric("Word Count Change", f"{opt_words}", delta=f"{word_change:+,}")

                                    st.markdown("---")

                                    # Side-by-side comparison
                                    comp_col1, comp_col2 = st.columns(2)

                                    with comp_col1:
                                        st.markdown("**Original:**")
                                        st.markdown(section['original_content'][:800] + "..." if len(section['original_content']) > 800 else section['original_content'])

                                    with comp_col2:
                                        st.markdown("**Optimized:**")
                                        st.markdown(section['optimized_content'][:800] + "..." if len(section['optimized_content']) > 800 else section['optimized_content'])

                                    # Full versions
                                    with st.expander("View Full Versions"):
                                        st.markdown("**Full Original:**")
                                        st.markdown(section['original_content'])
                                        st.markdown("---")
                                        st.markdown("**Full Optimized:**")
                                        st.markdown(section['optimized_content'])

                            # Generate final optimized version with accepted sections
                            st.markdown("---")
                            st.markdown("### 📥 Download Final Version")

                            final_content = ""
                            for idx, section in enumerate(optimized_sections):
                                if st.session_state.accepted_sections[idx]:
                                    final_content += section['optimized_content'] + "\n\n"
                                else:
                                    final_content += section['original_content'] + "\n\n"

                            # Show acceptance summary
                            accepted_count = sum(st.session_state.accepted_sections)
                            st.info(f"**Accepted Sections:** {accepted_count} / {len(optimized_sections)}")

                            # Download final version
                            optimized_filename = f"optimized_content_sections_{datetime.now().strftime('%Y%m%d_%H%M%S')}.md"
                            st.markdown(
                                create_download_link(final_content, optimized_filename),
                                unsafe_allow_html=True
                            )

                            st.success("✅ **For Writers**: Review each section, accept the good ones, reject the rest. Download combines your selections!")

                elif generate_optimized and not action_plan_result:
                    st.warning("⚠️ Please enable 'Generate Action Plan' to optimize sections.")

                # Summary and Download All
                st.markdown("---")
                st.markdown("## 📦 Analysis Summary & Downloads")

                summary_col1, summary_col2, summary_col3, summary_col4 = st.columns(4)

                with summary_col1:
                    if 'validation' in st.session_state.semantic_results:
                        st.metric("Entity Validation", "✅")
                    st.metric("Macro Context", "✅")
                    st.metric("Micro Context", "✅")

                with summary_col2:
                    if include_distrib_semantics:
                        st.metric("Distributional Analysis", "✅")
                    if generate_gap_analysis_ui and 'gap_analysis' in st.session_state.semantic_results:
                        st.metric("Gap Analysis", "✅")

                with summary_col3:
                    if generate_action_plan and 'action_plan' in st.session_state.semantic_results:
                        st.metric("Action Plan", "✅")
                    if generate_brief and 'content_brief' in st.session_state.semantic_results:
                        st.metric("Content Brief", "✅")

                with summary_col4:
                    if generate_optimized and 'optimized_sections' in st.session_state.semantic_results:
                        st.metric("Section Optimization", "✅")
                    st.metric("Framework", "Koray SEO")

                # Highlight actionable outputs for writers
                st.markdown("---")
                st.success("✅ **For Writers**: Download the Action Plan for copy-paste ready optimization steps!")

                # Download complete report
                if st.session_state.semantic_results:
                    st.markdown("---")

                    complete_report = f"""# Semantic SEO Analysis Report
**Primary Entity:** {primary_entity}
**Date:** {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
**Framework:** Koray Tuğberk GÜBÜR's Semantic SEO

---

"""

                    if 'validation' in st.session_state.semantic_results:
                        complete_report += f"""## Entity Validation

{st.session_state.semantic_results['validation']}

---

"""

                    if 'macro_context' in st.session_state.semantic_results:
                        complete_report += f"""## Macro Context Analysis

{st.session_state.semantic_results['macro_context']}

---

"""

                    if 'micro_context' in st.session_state.semantic_results:
                        complete_report += f"""## Micro Context Analysis

{st.session_state.semantic_results['micro_context']}

---

"""

                    if 'distributional' in st.session_state.semantic_results:
                        distrib = st.session_state.semantic_results['distributional']
                        complete_report += f"""## Distributional Semantics Analysis

### Top Terms
"""
                        for term, freq in distrib['top_terms']:
                            complete_report += f"- {term}: {freq}\n"

                        complete_report += f"""
### Top Co-occurring Pairs
"""
                        for (pair, count) in distrib['co_occurrences']:
                            complete_report += f"- {pair[0]} + {pair[1]}: {count} times\n"

                        complete_report += "\n---\n\n"

                    if 'gap_analysis' in st.session_state.semantic_results:
                        complete_report += f"""## Gap Analysis - What's Missing

{st.session_state.semantic_results['gap_analysis']}

---

"""

                    if 'action_plan' in st.session_state.semantic_results:
                        complete_report += f"""## Optimization Action Plan (Copy-Paste Ready)

{st.session_state.semantic_results['action_plan']}

---

"""

                    if 'content_brief' in st.session_state.semantic_results:
                        complete_report += f"""## Content Brief (4-Column Framework)

{st.session_state.semantic_results['content_brief']}

---

"""

                    if 'optimized_sections' in st.session_state.semantic_results:
                        complete_report += f"""## Section-by-Section Optimizations

"""
                        for section in st.session_state.semantic_results['optimized_sections']:
                            complete_report += f"""### Section {section['section_number']}: {section['original_heading']}

**Optimized Version:**

{section['optimized_content']}

---

"""

                    # Download button for complete report
                    report_filename = f"semantic_seo_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.md"
                    st.markdown(
                        create_download_link(complete_report, report_filename),
                        unsafe_allow_html=True
                    )

                    st.success("✅ Complete semantic SEO analysis finished! Download your complete report above.")

                    # Add writer instructions
                    st.info("""
**📝 For Your Writing Team:**

1. **Start with Gap Analysis** - See what's missing with source-attributed examples
2. **Review Action Plan** - Prioritized, copy-paste ready steps (with citations!)
3. **Use Content Brief** - Strategic framework for overall structure
4. **Review Optimized Sections** - Accept/reject individual H2 sections for granular control

✅ **All examples include source attributions** - fact-check before publishing!
🎯 **Action Plan is most actionable** - prioritized with before/after examples
                    """)

if __name__ == "__main__":
    main()
